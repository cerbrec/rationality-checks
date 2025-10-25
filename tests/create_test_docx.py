#!/usr/bin/env python3
"""
Script to create a test DOCX file with verifiable claims.

This creates a DOCX file similar to test_retzlaff.md but in DOCX format,
containing intentionally incorrect claims for testing verification.
"""

try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    print("❌ Error: python-docx not installed")
    print("Install with: pip install python-docx")
    exit(1)

from pathlib import Path

def create_test_docx():
    """Create a test DOCX file with claims that should be flagged"""

    # Create document
    doc = Document()

    # Title
    title = doc.add_heading('Test Athlete Profile - Sarah Johnson', 0)

    # Introduction
    doc.add_paragraph(
        "Sarah Johnson is a senior basketball player at State University, "
        "known for her exceptional three-point shooting and leadership on the court."
    )

    # Career Statistics (contains errors)
    doc.add_heading('Career Statistics', level=1)
    stats_para = doc.add_paragraph()
    stats_para.add_run("Season 2024-2025:\n").bold = True
    stats_para.add_run(
        "• Points per game: 28.5\n"  # Intentionally inflated
        "• Three-point percentage: 52%\n"  # Intentionally inflated
        "• Assists per game: 8.2\n"  # Intentionally inflated
        "• Rebounds per game: 9.1\n"
    )

    # Academic Information (contains errors)
    doc.add_heading('Academic Information', level=1)
    doc.add_paragraph(
        "Sarah is pursuing a degree in Mechanical Engineering with a perfect 4.0 GPA. "  # 4.0 is inflated
        "She has received the Dean's List honor for all eight semesters and was named "
        "Academic All-American in 2024."
    )

    # Social Media Following (contains errors)
    doc.add_heading('Social Media Presence', level=1)
    doc.add_paragraph(
        "Sarah has built a strong social media following:\n"
        "• Instagram: 95,000 followers\n"  # Intentionally inflated
        "• Twitter/X: 42,000 followers\n"  # Intentionally inflated
        "• TikTok: 180,000 followers"  # Intentionally inflated
    )

    # Awards (mix of correct and questionable claims)
    doc.add_heading('Awards and Recognition', level=1)
    doc.add_paragraph(
        "• 2024 Conference Player of the Year\n"
        "• 2023 All-Conference First Team\n"
        "• 2024 State Player of the Year\n"
        "• Named to the Naismith Trophy Watch List"
    )

    # Background (contains a factual error)
    doc.add_heading('Background', level=1)
    doc.add_paragraph(
        "Originally from Seattle, Washington, Sarah led her high school team to three "
        "consecutive state championships from 2018-2020. She was a five-star recruit "  # Intentionally inflated from typical 4-star
        "and ranked #8 in her recruiting class nationally."
    )

    # Future Plans
    doc.add_heading('Professional Prospects', level=1)
    doc.add_paragraph(
        "Sarah is projected as a first-round pick in the upcoming WNBA Draft, "
        "with several teams expressing strong interest in her skills as a combo guard."
    )

    # Save document
    output_path = Path(__file__).parent / "test_sarah_johnson.docx"
    doc.save(str(output_path))

    print(f"✓ Test DOCX file created: {output_path}")
    print(f"\nExpected issues to be flagged:")
    print("  • Points per game: 28.5 (unrealistically high for college basketball)")
    print("  • Three-point percentage: 52% (unrealistically high)")
    print("  • Assists per game: 8.2 (unrealistically high for a shooting guard)")
    print("  • 4.0 GPA in Mechanical Engineering (unlikely to verify)")
    print("  • Instagram: 95,000 followers (likely inflated)")
    print("  • Twitter: 42,000 followers (likely inflated)")
    print("  • TikTok: 180,000 followers (likely inflated)")
    print("  • Five-star recruit ranking (likely inflated)")

    return output_path


if __name__ == "__main__":
    create_test_docx()
